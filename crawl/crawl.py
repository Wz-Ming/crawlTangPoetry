import requests
from pathlib import Path
from bs4 import BeautifulSoup
import bs4
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from ebooklib import epub
import re

# def saveToTxt(contents:list):
#     fileName='唐诗三百首.txt'
#     with open(fileName,'a',encoding='utf-8') as f:
#         print('写入中...'+contents[0])
#         for rowContent in contents:  
#             f.write(rowContent+'\n')

def getHtmlText(url:str)->str:
    try:
        simBrowser={
            'User-Agent':'Chrome/114.0.5735.106',
            'Accept-Encoding': 'gzip, deflate', 
            'Accept': '*/*', 
            'Connection': 'keep-alive'
        }
        requests.url
        r=requests.get(url,headers=simBrowser,timeout=30)
        r.raise_for_status()
        r.encoding=r.apparent_encoding
        return r.text
    except:
        return str()

def getPoetryContent(soup:BeautifulSoup,contentList:list,string:str):
    for br_tag in soup.find_all('br'):
        br_tag.replace_with('\n')
    for tag in soup.find_all(id='contson'+string):
        content=tag.text.strip()
        content=re.sub(r'(\(|（).*?(\)|）)','',content)
        contentList.append(content)
    return contentList

def clearDocx(fileName:str):
    path=Path(fileName)
    if path.exists():
        doc=Document(fileName)
        for element in doc.element.body:
            doc.element.body.remove(element)
        doc.save(fileName)

def addEmptyLine(doc:Document):
    text=doc.add_paragraph()
    format=text.paragraph_format
    #format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    format.line_spacing = 1.0  # 设置为1.0表示单倍行距
    format.space_before = 0
    format.space_after = format.space_before

def paragraphCommonFormat(run,format):
    run.font.name = 'SimSun'
    run.font.color.rgb = RGBColor(0, 0, 0) 
    format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    format.line_spacing = 1.0  # 设置为1.0表示单倍行距
    format.space_before = 0
    format.space_after = format.space_before

def setDocxTitle(fileName:str,string:str,rank:int):
    path=Path(fileName)
    doc:Document=Document()
    if path.exists():
        doc=Document(fileName)
    title=doc.add_paragraph(string,'Heading'+str(rank))
    titleRun=title.runs[0]
    titleParagraph_format = title.paragraph_format
    paragraphCommonFormat(titleRun,titleParagraph_format)
    #addEmptyLine(doc)
    doc.save(fileName)

def writeToDocx(fileName:str,contents:list):
    path=Path(fileName)
    doc:Document=Document()
    if path.exists():
        doc=Document(fileName)
    pos=0
    for content in contents:
        paragraph=None
        if pos==0:
            paragraph=doc.add_paragraph(content,'Heading3')
        # elif pos==1:
        #     paragraph=doc.add_paragraph(content,'Heading4')
        else:
            paragraph=doc.add_paragraph(content)
        textRun=paragraph.runs[0]
        Paragraph_format = paragraph.paragraph_format
        paragraphCommonFormat(textRun,Paragraph_format)
        pos+=1
    #addEmptyLine(doc)
    doc.save(fileName)

def saveToDocx(htmlText:str):
    bookName='唐诗三百首'
    saveFileName='/home/linux/Documents/books/{}.docx'.format(bookName)
    clearDocx(saveFileName)
    #setDocxTitle(saveFileName,bookName,1)
    subTitlePos=0
    regex = re.compile(r'(\(|（).*?(\)|）)')
    for div in BeautifulSoup(htmlText,'html.parser').find_all(attrs='typecont'):
        subTitlePos+=1
        setDocxTitle(saveFileName,str(subTitlePos)+'、'+div.div.text,1)
        for tag in div(name='span'):
            poetryName=tag.text.strip()
            author=regex.search(poetryName)
            poetryName=regex.sub('',poetryName)
            contentList=[]
            contentList.append(poetryName)
            if(author):
                contentList.append(author.group(0))
            newUrl='https://so.gushiwen.cn'+tag.a.get('href')
            getPoetryContent(BeautifulSoup(getHtmlText(newUrl),'html.parser'),contentList,newUrl[-17:-5])
            writeToDocx(saveFileName,contentList)
            print('爬取中...'+tag.text.strip())

def saveToEpub(htmlText:str):
    bookName='唐诗三百首'
    saveFileName='/home/linux/Documents/books/{}.epub'.format(bookName)
    book = epub.EpubBook()
    book.set_title(bookName)
    book.set_language('zh')
    subTitlePos=0
    regex = re.compile(r'(\(|（).*?(\)|）)')
    for div in BeautifulSoup(htmlText,'html.parser').find_all(attrs='typecont'):
        subTitlePos+=1
        mainTitle=str(subTitlePos)+'、'+div.div.text
        mainHtmlName='{}.xhtml'.format(mainTitle)
        mainHtmlContent="<h1>{}</h1>".format(mainTitle)
        mainChapter=epub.EpubHtml(title=mainTitle,file_name=mainHtmlName,content=mainHtmlContent)
        book.add_item(mainChapter)
        chapterList:list[epub.EpubHtml]=[]
        for titleTag in div(name='span'):
            subTitle=titleTag.text.strip()
            poetryName=subTitle
            author=regex.search(poetryName)
            poetryName=regex.sub('',subTitle)
            subHtmlName='{}.xhtml'.format(subTitle)
            suHtmlContent="<h2>{}</h2>".format(poetryName)
            suHtmlContent+='<style>body{text-align:center;}</style><p>'
            if(author):
                suHtmlContent+=author.group(0)+'<br>'
            newUrl='https://so.gushiwen.cn'+titleTag.a.get('href')
            for contentTag in BeautifulSoup(getHtmlText(newUrl),'html.parser').find_all(id='contson'+newUrl[-17:-5]):
                for br_tag in contentTag.find_all('br'):
                    br_tag.replace_with('@')
                content=re.sub(r'@', "<br>",contentTag.text.strip())
                content=regex.sub('',content)
                suHtmlContent+=content
            suHtmlContent+='</p>'
            subChapter=epub.EpubHtml(title=subTitle,file_name=subHtmlName,content=suHtmlContent)
            book.add_item(subChapter)
            chapterList.append(subChapter)
            print('爬取中...'+subTitle)
        book.toc.append((epub.Section(mainTitle,mainHtmlName),tuple(chapterList)))
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
    epub.write_epub(saveFileName, book,{})    

def CrawlTangPoetry(htmlText:str):
    #saveToDocx(htmlText)
    saveToEpub(htmlText)

def main():
    print('开始爬取！')
    url='https://so.gushiwen.cn/gushi/tangshi.aspx'
    htmlText=getHtmlText(url)
    if len(htmlText)>0:
        print('爬取成功！')
        CrawlTangPoetry(htmlText)
    else:
        print('爬取失败！')

main()